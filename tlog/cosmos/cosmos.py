from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import psycopg2
from datetime import datetime

DB_CONFIG = {
    "host": "192.168.31.5",
    "port": 5432,
    "database": "tlog",
    "user": "tlog",
    "password": "Labrador603502$"
}

def format_date(date_val):
    if date_val is None:
        return ""
    if hasattr(date_val, 'strftime'):
        return date_val.strftime("%d.%m.%Y")
    return str(date_val)

def format_time(time_val):
    if time_val is None:
        return ""
    if hasattr(time_val, 'strftime'):
        return time_val.strftime("%H:%M")
    return str(time_val)

try:
    conn = psycopg2.connect(**DB_CONFIG)
    cursor = conn.cursor()
    cursor.execute("SELECT 1")
    result = cursor.fetchone()
    if result:
        print("Подключение выполнено")
    conn.close()
except Exception as e:
    print(f"Ошибка подключения: {e}")
    exit(1)

doc_path = "cosmos.docx"
name = input("Введите имя и фамилию: ")
my_callsign = input("Введите ваш позывной: ").upper()
other_callsigns = input("Введите другие позывные: ").upper()
your_email = input("Введите ваш email: ")
your_phone = input("Введите ваш телефон: ")
info = input("Введите информацию: ")

doc = Document(doc_path)

qso_data = []
try:
    conn = psycopg2.connect(**DB_CONFIG)
    cursor = conn.cursor()
    cursor.execute("""
        SELECT * FROM (
            SELECT DISTINCT ON (callsign)
                LEFT(COALESCE(gridsquare, ''), 4) as gridsquare,
                date, band, time, mode, callsign,
                rst_sent, rst_rcvd
            FROM tlog_qso
            WHERE my_callsign = %s AND prop_mode = 'SAT'
            ORDER BY callsign, date, time
        ) AS distinct_qsos
        ORDER BY date, time
    """, (my_callsign,))
    qso_data = cursor.fetchall()
    conn.close()
    print(f"Найдено {len(qso_data)} уникальных записей QSO")
except Exception as e:
    print(f"Ошибка запроса к БД: {e}")
    exit(1)

if len(qso_data) < 100:
    print("Недостаточно позывных")
    exit(1)

count25 = (len(qso_data) // 25) * 25
current_date = datetime.now().strftime("%d.%m.%Y")
first_word_name = name.split()[0] if name else ""

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml("<w:tcBorders {}><w:top w:val=\"single\" w:sz=\"2\"/><w:left w:val=\"single\" w:sz=\"2\"/><w:right w:val=\"single\" w:sz=\"2\"/><w:bottom w:val=\"single\" w:sz=\"2\"/><w:insideH w:val=\"single\" w:sz=\"2\"/><w:insideV w:val=\"single\" w:sz=\"2\"/></w:tcBorders>".format(nsdecls("w")))
    tcPr.append(borders)

def set_cell_shading(cell, color):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_col_width(table, col_idx, width):
    for row in table.rows:
        row.cells[col_idx].width = width

for paragraph in doc.paragraphs:
    if "{{YOUR_NAME}}" in paragraph.text:
        parts = paragraph.text.split("{{YOUR_NAME}}")
        paragraph.clear()
        for i, part in enumerate(parts[:-1]):
            paragraph.add_run(part)
            run = paragraph.add_run(name)
            run.bold = True
            run.font.size = Pt(14)
        paragraph.add_run(parts[-1])
    elif "{{MY_CALLSIGN}}" in paragraph.text:
        parts = paragraph.text.split("{{MY_CALLSIGN}}")
        paragraph.clear()
        for i, part in enumerate(parts[:-1]):
            paragraph.add_run(part)
            run = paragraph.add_run(my_callsign)
            run.bold = True
            run.font.size = Pt(14)
        paragraph.add_run(parts[-1])
    elif "{{OTHER_CALLSIGNS}}" in paragraph.text:
        parts = paragraph.text.split("{{OTHER_CALLSIGNS}}")
        paragraph.clear()
        for i, part in enumerate(parts[:-1]):
            paragraph.add_run(part)
            run = paragraph.add_run(other_callsigns)
            run.bold = True
            run.font.size = Pt(14)
        paragraph.add_run(parts[-1])
    elif "{{YOUR_EMAIL}}" in paragraph.text:
        parts = paragraph.text.split("{{YOUR_EMAIL}}")
        paragraph.clear()
        for i, part in enumerate(parts[:-1]):
            paragraph.add_run(part)
            run = paragraph.add_run(your_email)
            run.bold = True
            run.font.size = Pt(14)
        paragraph.add_run(parts[-1])
    elif "{{YOUR_PHONE}}" in paragraph.text:
        parts = paragraph.text.split("{{YOUR_PHONE}}")
        paragraph.clear()
        for i, part in enumerate(parts[:-1]):
            paragraph.add_run(part)
            run = paragraph.add_run(your_phone)
            run.bold = True
            run.font.size = Pt(14)
        paragraph.add_run(parts[-1])
    elif "{{INITIAL}}" in paragraph.text:
        parts = paragraph.text.split("{{INITIAL}}")
        paragraph.clear()
        for i, part in enumerate(parts[:-1]):
            paragraph.add_run(part)
            run = paragraph.add_run(str(len(qso_data)))
            run.bold = True
            run.font.size = Pt(14)
        paragraph.add_run(parts[-1])
    elif "{{INFO}}" in paragraph.text:
        parts = paragraph.text.split("{{INFO}}")
        paragraph.clear()
        for i, part in enumerate(parts[:-1]):
            paragraph.add_run(part)
            run = paragraph.add_run(info)
            run.bold = True
            run.font.size = Pt(14)
        paragraph.add_run(parts[-1])
    elif "{{BOTTOM_LINE}}" in paragraph.text:
        paragraph.clear()
        run = paragraph.add_run("Подпись ")
        run.font.size = Pt(14)
        run = paragraph.add_run(first_word_name)
        run.bold = True
        run.font.size = Pt(14)
        run = paragraph.add_run("       Позывной ")
        run.font.size = Pt(14)
        run = paragraph.add_run(my_callsign)
        run.bold = True
        run.font.size = Pt(14)
        run = paragraph.add_run("          Дата ")
        run.font.size = Pt(14)
        run = paragraph.add_run(current_date)
        run.bold = True
        run.font.size = Pt(14)
    elif "{{count25}}" in paragraph.text:
        original_run = paragraph.runs[0] if paragraph.runs else None
        original_font_size = original_run.font.size if original_run else None
        original_font_name = original_run.font.name if original_run else None
        original_bold = original_run.bold if original_run else None

        parts = paragraph.text.split("{{count25}}")
        paragraph.clear()
        for i, part in enumerate(parts[:-1]):
            paragraph.add_run(part)
            if original_font_size:
                paragraph.runs[-1].font.size = original_font_size
            if original_font_name:
                paragraph.runs[-1].font.name = original_font_name
            if original_bold:
                paragraph.runs[-1].bold = True
        run = paragraph.add_run(str(count25))
        run.bold = True
        run.font.size = Pt(22)
        paragraph.add_run(parts[-1])
        if original_font_size:
            paragraph.runs[-1].font.size = original_font_size
        if original_font_name:
            paragraph.runs[-1].font.name = original_font_name
        if original_bold:
            paragraph.runs[-1].bold = True
    elif "{{TABLE}}" in paragraph.text:
        paragraph.clear()
        if qso_data:
            table = doc.add_table(rows=1, cols=9)
            headers = ["№", "QTH-loc", "Дата", "Диапазон", "Время", "Модуляция", "Позывной", "Рапорт переданный", "Рапорт принятый"]
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                run = header_cells[i].paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(10)
                set_cell_border(header_cells[i])
                set_cell_shading(header_cells[i], "D9D9D9")
            for row_num, row_data in enumerate(qso_data, 1):
                row = table.add_row().cells
                row[0].text = str(row_num)
                row[1].text = row_data[0] or ""
                row[2].text = format_date(row_data[1])
                row[3].text = row_data[2] or ""
                row[4].text = format_time(row_data[3])
                row[5].text = row_data[4] or ""
                row[6].text = row_data[5] or ""
                row[7].text = row_data[6] or ""
                row[8].text = row_data[7] or ""
                for cell in row:
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
                    set_cell_border(cell)
            set_col_width(table, 0, Pt(10))

doc.save("CosmosAplay.docx")
print("\nГотово! Документ сохранён как CosmosAplay.docx")
