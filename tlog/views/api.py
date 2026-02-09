# API функции

from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from django.db import models


@login_required
def api_user_info(request):
    """
    API для получения информации о пользователе
    """
    try:
        from ..models import RadioProfile, QSO
        
        # Базовая информация о пользователе
        user_info = {
            'id': request.user.id,
            'username': request.user.username,
            'email': request.user.email,
            'first_name': request.user.first_name,
            'last_name': request.user.last_name,
            'date_joined': request.user.date_joined.isoformat(),
        }
        
        # Информация из профиля радиолюбителя
        try:
            profile = request.user.radio_profile
            user_info.update({
                'callsign': profile.callsign,
                'qth': profile.qth,
                'my_gridsquare': profile.my_gridsquare,
                'lotw_user': profile.lotw_user,
                'lotw_chk_pass': profile.lotw_chk_pass,
            })
        except RadioProfile.DoesNotExist:
            user_info['callsign'] = request.user.username
        
        # Статистика QSO
        user_qso = QSO.objects.filter(user=request.user)
        user_info['qso_stats'] = {
            'total_qso': user_qso.count(),
            'unique_callsigns': user_qso.values('callsign').distinct().count(),
            'dxcc_count': user_qso.exclude(dxcc__isnull=True).exclude(dxcc='').values('dxcc').distinct().count(),
        }
        
        return JsonResponse({
            'success': True,
            'user_info': user_info
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


@login_required
def api_qso_stats(request):
    """
    API для получения статистики QSO пользователя
    """
    try:
        from ..models import QSO
        from collections import Counter
        
        user_qso = QSO.objects.filter(user=request.user)
        
        # Общая статистика
        total_qso = user_qso.count()
        unique_callsigns = user_qso.values('callsign').distinct().count()
        
        # Статистика по диапазонам
        bands = user_qso.values_list('band', flat=True)
        band_stats = dict(Counter(bands))
        
        # Статистика по модуляции
        modes = user_qso.values_list('mode', flat=True)
        mode_stats = dict(Counter(modes))
        
        # Статистика по годам
        years = [qso.date.year for qso in user_qso]
        year_stats = dict(Counter(years))
        
        # Последние QSO
        recent_qso = user_qso.order_by('-date', '-time')[:10].values(
            'callsign', 'date', 'time', 'band', 'mode', 'gridsquare'
        )
        
        return JsonResponse({
            'success': True,
            'stats': {
                'total_qso': total_qso,
                'unique_callsigns': unique_callsigns,
                'band_statistics': band_stats,
                'mode_statistics': mode_stats,
                'year_statistics': year_stats,
                'recent_qso': list(recent_qso),
            }
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


@login_required
def api_search_callsigns(request):
    """
    API для поиска позывных (аналог get_callsigns_list)
    """
    query = request.GET.get('q', '').upper()
    limit = int(request.GET.get('limit', 10))
    
    if len(query) < 1:
        return JsonResponse({'callsigns': []})
    
    try:
        from ..models import QSO, RadioProfile
        import json
        
        callsigns_set = set()
        
        # 1. Ищем в поле callsign модели RadioProfile
        profiles_with_callsign = RadioProfile.objects.filter(
            callsign__icontains=query
        ).exclude(callsign='').values_list('callsign', flat=True).distinct()
        callsigns_set.update([c.upper() for c in profiles_with_callsign])
        
        # 2. Ищем в поле my_callsigns (JSON) модели RadioProfile
        profiles = RadioProfile.objects.filter(
            my_callsigns__isnull=False
        ).exclude(
            my_callsigns=[]
        ).exclude(
            my_callsigns=''
        )
        
        for profile in profiles:
            try:
                my_callsigns = json.loads(profile.my_callsigns) if isinstance(profile.my_callsigns, str) else profile.my_callsigns
                if isinstance(my_callsigns, list):
                    for item in my_callsigns:
                        if isinstance(item, dict) and 'name' in item:
                            name = item['name'].upper()
                            if query in name:
                                callsigns_set.add(name)
                        elif isinstance(item, str):
                            name = item.upper()
                            if query in name:
                                callsigns_set.add(name)
            except (json.JSONDecodeError, TypeError):
                pass
        
        # 3. Также ищем в QSO (my_callsign)
        qso_callsigns = QSO.objects.filter(
            my_callsign__icontains=query
        ).values_list('my_callsign', flat=True).distinct()
        callsigns_set.update([c.upper() for c in qso_callsigns])
        
        # Ограничиваем и сортируем
        callsigns_list = sorted(list(callsigns_set))[:limit]
        
        return JsonResponse({
            'success': True,
            'callsigns': callsigns_list
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


@login_required
def api_cosmos_user_data(request):
    """
    API для получения данных пользователя для формы Cosmos
    """
    try:
        from ..models import RadioProfile
        import json

        user = request.user
        profile = None
        try:
            profile = user.radio_profile
        except RadioProfile.DoesNotExist:
            pass

        # Формируем ФИО из базы
        full_name = ''
        if profile:
            if profile.first_name and profile.last_name:
                full_name = f"{profile.last_name} {profile.first_name}"
            elif profile.last_name:
                full_name = profile.last_name
            elif profile.first_name:
                full_name = profile.first_name
        if not full_name and (user.first_name or user.last_name):
            full_name = f"{user.last_name or ''} {user.first_name or ''}".strip()

        # Получаем email из базы
        email = user.email or ''

        # Получаем основной позывной (user.username)
        main_callsign = user.username.upper()

        # Получаем дополнительные позывные из профиля
        other_callsigns_list = []
        if profile and profile.my_callsigns:
            try:
                my_callsigns = profile.my_callsigns
                if isinstance(my_callsigns, str):
                    my_callsigns = json.loads(my_callsigns)
                if isinstance(my_callsigns, list):
                    for item in my_callsigns:
                        if isinstance(item, dict):
                            if item.get('name'):
                                other_callsigns_list.append(item['name'].upper())
                        elif isinstance(item, str):
                            other_callsigns_list.append(item.upper())
            except (json.JSONDecodeError, TypeError):
                pass

        # Удаляем основной позывной из списка дополнительных, если он там есть
        other_callsigns_list = [c for c in other_callsigns_list if c != main_callsign]

        return JsonResponse({
            'success': True,
            'main_callsign': main_callsign,
            'full_name': full_name,
            'email': email,
            'phone': '',
            'info': '',
            'other_callsigns': other_callsigns_list,
        })

    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


@login_required
def api_cosmos_generate(request):
    """
    API для генерации заявки на диплом Cosmos
    """
    try:
        from ..models import RadioProfile, QSO
        from django.conf import settings
        import os
        import json
        import time
        from datetime import datetime
        from io import BytesIO

        # Получаем данные из формы
        main_callsign = request.POST.get('main_callsign', '').strip().upper()
        full_name = request.POST.get('full_name', '').strip()
        email = request.POST.get('email', '').strip()
        phone = request.POST.get('phone', '').strip()
        info = request.POST.get('info', '').strip()

        # Получаем дополнительные позывные
        other_callsigns_raw = request.POST.get('other_callsigns_json', '[]')
        try:
            other_callsigns = json.loads(other_callsigns_raw)
            if isinstance(other_callsigns, list):
                other_callsigns = [c.upper() if isinstance(c, str) else c for c in other_callsigns]
        except (json.JSONDecodeError, TypeError):
            other_callsigns = []

        # Валидация
        if not main_callsign:
            return JsonResponse({
                'success': False,
                'error': 'Позывной обязателен'
            }, status=400)

        if not full_name:
            return JsonResponse({
                'success': False,
                'error': 'ФИО обязательно'
            }, status=400)

        if not email:
            return JsonResponse({
                'success': False,
                'error': 'Email обязателен'
            }, status=400)

        # Формируем список всех позывных для поиска QSO
        all_callsigns = [main_callsign] + [c for c in other_callsigns if c and c != main_callsign]

        # Оптимизированный запрос QSO через ORM с использованием индексов
        qso_data = []
        try:
            if all_callsigns:
                # Используем ORM для оптимизированного запроса
                qsos = QSO.objects.filter(
                    user=request.user,
                    my_callsign__in=all_callsigns
                ).filter(
                    models.Q(prop_mode='SAT') | models.Q(band='13CM')
                ).order_by('callsign', 'date', 'time')

                # Получаем только нужные поля
                qsos = qsos.values('callsign', 'date', 'time', 'band', 'mode', 'rst_sent', 'rst_rcvd', 'gridsquare')

                # Убираем дубликаты по callsign (берем первую запись)
                seen_callsigns = set()
                for qso in qsos:
                    if qso['callsign'] not in seen_callsigns:
                        seen_callsigns.add(qso['callsign'])
                        # Формируем gridsquare (первые 4 символа)
                        gridsquare = qso.get('gridsquare', '') or ''
                        gridsquare_4 = gridsquare[:4] if len(gridsquare) >= 4 else gridsquare
                        qso_data.append((
                            gridsquare_4,
                            qso['date'],
                            qso['band'],
                            qso['time'],
                            qso['mode'],
                            qso['callsign'],
                            qso['rst_sent'],
                            qso['rst_rcvd']
                        ))

                # Сортируем по дате и времени
                qso_data.sort(key=lambda x: (x[1] or '', x[3] or ''))
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': f'Не удалось получить данные QSO: {str(e)}'
            }, status=500)

        # Путь к шаблону
        template_path = os.path.join(settings.BASE_DIR, 'tlog', 'cosmos', 'Cosmos.docx')

        if not os.path.exists(template_path):
            return JsonResponse({
                'success': False,
                'error': 'Шаблон документа не найден'
            }, status=500)

        # Открываем документ
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls

            doc = Document(template_path)
        except ImportError:
            return JsonResponse({
                'success': False,
                'error': 'Библиотека python-docx не установлена'
            }, status=500)

        # Форматируем данные
        def format_date(date_val):
            if date_val is None:
                return ""
            if hasattr(date_val, 'strftime'):
                return date_val.strftime("%d.%m.%Y")
            return str(date_val)

        def format_time(time_val):
            if time_val is None:
                return ""
            if hasattr(time_val, 'strftime'):
                return time_val.strftime("%H:%M")
            return str(time_val)

        def set_cell_border(cell):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                "<w:tcBorders {}><w:top w:val=\"single\" w:sz=\"2\"/><w:left w:val=\"single\" w:sz=\"2\"/><w:right w:val=\"single\" w:sz=\"2\"/><w:bottom w:val=\"single\" w:sz=\"2\"/><w:insideH w:val=\"single\" w:sz=\"2\"/><w:insideV w:val=\"single\" w:sz=\"2\"/></w:tcBorders>".format(
                    nsdecls("w")))
            tcPr.append(borders)

        def set_cell_shading(cell, color):
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
            cell._tc.get_or_add_tcPr().append(shading_elm)

        def set_col_width(table, col_idx, width):
            for row in table.rows:
                row.cells[col_idx].width = width

        # Формируем строку с другими позывными
        other_callsigns_str = ', '.join(other_callsigns) if other_callsigns else ''

        # Подсчеты
        count25 = (len(qso_data) // 25) * 25
        current_date = datetime.now().strftime("%d.%m.%Y")
        first_word_name = full_name.split()[0] if full_name else ""

        # Заполняем документ
        for paragraph in doc.paragraphs:
            if "{{YOUR_NAME}}" in paragraph.text:
                parts = paragraph.text.split("{{YOUR_NAME}}")
                paragraph.clear()
                for part in parts[:-1]:
                    paragraph.add_run(part)
                    run = paragraph.add_run(full_name)
                    run.bold = True
                    run.font.size = Pt(14)
                paragraph.add_run(parts[-1])
            elif "{{MY_CALLSIGN}}" in paragraph.text:
                parts = paragraph.text.split("{{MY_CALLSIGN}}")
                paragraph.clear()
                for part in parts[:-1]:
                    paragraph.add_run(part)
                    run = paragraph.add_run(main_callsign)
                    run.bold = True
                    run.font.size = Pt(14)
                paragraph.add_run(parts[-1])
            elif "{{OTHER_CALLSIGNS}}" in paragraph.text:
                parts = paragraph.text.split("{{OTHER_CALLSIGNS}}")
                paragraph.clear()
                for part in parts[:-1]:
                    paragraph.add_run(part)
                    run = paragraph.add_run(other_callsigns_str)
                    run.bold = True
                    run.font.size = Pt(14)
                paragraph.add_run(parts[-1])
            elif "{{YOUR_EMAIL}}" in paragraph.text:
                parts = paragraph.text.split("{{YOUR_EMAIL}}")
                paragraph.clear()
                for part in parts[:-1]:
                    paragraph.add_run(part)
                    run = paragraph.add_run(email)
                    run.bold = True
                    run.font.size = Pt(14)
                paragraph.add_run(parts[-1])
            elif "{{YOUR_PHONE}}" in paragraph.text:
                parts = paragraph.text.split("{{YOUR_PHONE}}")
                paragraph.clear()
                for part in parts[:-1]:
                    paragraph.add_run(part)
                    run = paragraph.add_run(phone)
                    run.bold = True
                    run.font.size = Pt(14)
                paragraph.add_run(parts[-1])
            elif "{{INITIAL}}" in paragraph.text:
                parts = paragraph.text.split("{{INITIAL}}")
                paragraph.clear()
                for part in parts[:-1]:
                    paragraph.add_run(part)
                    run = paragraph.add_run(str(len(qso_data)))
                    run.bold = True
                    run.font.size = Pt(14)
                paragraph.add_run(parts[-1])
            elif "{{INFO}}" in paragraph.text:
                parts = paragraph.text.split("{{INFO}}")
                paragraph.clear()
                for part in parts[:-1]:
                    paragraph.add_run(part)
                    run = paragraph.add_run(info)
                    run.bold = True
                    run.font.size = Pt(14)
                paragraph.add_run(parts[-1])
            elif "{{BOTTOM_LINE}}" in paragraph.text:
                paragraph.clear()
                run = paragraph.add_run("Подпись ")
                run.font.size = Pt(14)
                run = paragraph.add_run(first_word_name)
                run.bold = True
                run.font.size = Pt(14)
                run = paragraph.add_run("       Позывной ")
                run.font.size = Pt(14)
                run = paragraph.add_run(main_callsign)
                run.bold = True
                run.font.size = Pt(14)
                run = paragraph.add_run("          Дата ")
                run.font.size = Pt(14)
                run = paragraph.add_run(current_date)
                run.bold = True
                run.font.size = Pt(14)
            elif "{{count25}}" in paragraph.text:
                original_run = paragraph.runs[0] if paragraph.runs else None
                original_font_size = original_run.font.size if original_run else None
                original_font_name = original_run.font.name if original_run else None
                original_bold = original_run.bold if original_run else None

                parts = paragraph.text.split("{{count25}}")
                paragraph.clear()
                for part in parts[:-1]:
                    paragraph.add_run(part)
                    if original_font_size:
                        paragraph.runs[-1].font.size = original_font_size
                    if original_font_name:
                        paragraph.runs[-1].font.name = original_font_name
                    if original_bold:
                        paragraph.runs[-1].bold = True
                run = paragraph.add_run(str(count25))
                run.bold = True
                run.font.size = Pt(22)
                paragraph.add_run(parts[-1])
                if original_font_size:
                    paragraph.runs[-1].font.size = original_font_size
                if original_font_name:
                    paragraph.runs[-1].font.name = original_font_name
                if original_bold:
                    paragraph.runs[-1].bold = True
            elif "{{TABLE}}" in paragraph.text:
                paragraph.clear()
                if qso_data:
                    table = doc.add_table(rows=1, cols=9)
                    headers = ["№", "QTH-loc", "Дата", "Диапазон", "Время", "Модуляция", "Позывной", "Рапорт переданный", "Рапорт принятый"]
                    header_cells = table.rows[0].cells
                    for i, header in enumerate(headers):
                        header_cells[i].text = header
                        run = header_cells[i].paragraphs[0].runs[0]
                        run.bold = True
                        run.font.size = Pt(10)
                        set_cell_border(header_cells[i])
                        set_cell_shading(header_cells[i], "D9D9D9")
                    for row_num, row_data in enumerate(qso_data, 1):
                        row = table.add_row().cells
                        row[0].text = str(row_num)
                        row[1].text = row_data[0] or ""
                        row[2].text = format_date(row_data[1])
                        row[3].text = row_data[2] or ""
                        row[4].text = format_time(row_data[3])
                        row[5].text = row_data[4] or ""
                        row[6].text = row_data[5] or ""
                        row[7].text = row_data[6] or ""
                        row[8].text = row_data[7] or ""
                        for cell in row:
                            cell.paragraphs[0].runs[0].font.size = Pt(10)
                            set_cell_border(cell)
                    set_col_width(table, 0, Pt(10))

        # Формируем имя файла
        safe_username = request.user.username.replace(' ', '_')
        output_filename = f"Заявка_Cosmos_{safe_username}.docx"

        # Сохраняем во временный буфер
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Сохраняем файл во временный файл для скачивания
        temp_dir = os.path.join(settings.BASE_DIR, 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        temp_file_path = os.path.join(temp_dir, f'cosmos_{request.user.id}_{int(time.time())}.docx')

        with open(temp_file_path, 'wb') as f:
            f.write(buffer.getvalue())

        # Сохраняем путь к файлу в сессии
        request.session['cosmos_download_path'] = temp_file_path
        request.session['cosmos_download_filename'] = output_filename

        # Формируем сообщение
        if len(qso_data) >= 100:
            message = f'✅ Заявка сформирована! Найдено {len(qso_data)} уникальных QSO.'
        else:
            message = f'⚠️ Внимание! В заявке найдено только {len(qso_data)} уникальных QSO. Для диплома необходимо минимум 100.'

        return JsonResponse({
            'success': True,
            'message': message,
            'qso_count': len(qso_data)
        })

    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f'Ошибка при формировании заявки: {str(e)}'
        }, status=500)


@login_required
def api_cosmos_download(request):
    """
    API для скачивания сформированного файла заявки Cosmos
    """
    try:
        import os
        from django.utils.encoding import escape_uri_path

        download_path = request.session.get('cosmos_download_path')
        download_filename = request.session.get('cosmos_download_filename')

        if not download_path or not os.path.exists(download_path):
            return JsonResponse({
                'success': False,
                'error': 'Файл заявки не найден. Пожалуйста, сформируйте заявку заново.'
            }, status=404)

        with open(download_path, 'rb') as f:
            content = f.read()

        # Удаляем временный файл
        os.remove(download_path)

        # Очищаем сессию
        if 'cosmos_download_path' in request.session:
            del request.session['cosmos_download_path']
        if 'cosmos_download_filename' in request.session:
            del request.session['cosmos_download_filename']

        from django.http import HttpResponse
        response = HttpResponse(
            content,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{escape_uri_path(download_filename)}"'
        return response

    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f'Ошибка при скачивании файла: {str(e)}'
        }, status=500)