# Функции диплома Cosmos

from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import connection
from django.conf import settings
import os
import json
import time
from datetime import datetime
from ..models import RadioProfile, check_user_blocked


@login_required
def cosmos_diploma(request):
    """
    Страница заявки на диплом Космос
    """
    if not request.user.is_authenticated:
        return redirect('login_page')

    # Проверяем, не заблокирован ли пользователь
    is_blocked, reason = check_user_blocked(request.user)
    if is_blocked:
        return render(request, 'blocked.html', {'reason': reason})

    # Получаем данные пользователя из базы
    user = request.user
    profile = None
    try:
        profile = user.radio_profile
    except RadioProfile.DoesNotExist:
        pass

    # Формируем ФИО из базы
    full_name = ''
    if profile:
        if profile.first_name and profile.last_name:
            full_name = f"{profile.last_name} {profile.first_name}"
        elif profile.last_name:
            full_name = profile.last_name
        elif profile.first_name:
            full_name = profile.first_name
    if not full_name and (user.first_name or user.last_name):
        full_name = f"{user.last_name or ''} {user.first_name or ''}".strip()

    # Получаем email из базы
    email = user.email or ''

    # Получаем основной позывной (user.username)
    main_callsign = user.username.upper()

    # Получаем дополнительные позывные из профиля
    other_callsigns_list = []
    if profile and profile.my_callsigns:
        try:
            my_callsigns = profile.my_callsigns
            if isinstance(my_callsigns, str):
                my_callsigns = json.loads(my_callsigns)
            if isinstance(my_callsigns, list):
                for item in my_callsigns:
                    if isinstance(item, dict):
                        if item.get('name'):
                            other_callsigns_list.append(item['name'].upper())
                    elif isinstance(item, str):
                        other_callsigns_list.append(item.upper())
        except (json.JSONDecodeError, TypeError):
            pass

    # Удаляем основной позывной из списка дополнительных, если он там есть
    other_callsigns_list = [c for c in other_callsigns_list if c != main_callsign]

    # Загружаем данные из куки для предзаполнения формы
    def get_cookie(name, default=''):
        return request.COOKIES.get(name, default) or default

    # Инициализируем данные формы
    form_data = {
        'main_callsign': get_cookie('cosmos_main_callsign', main_callsign),
        'full_name': get_cookie('cosmos_full_name', full_name),
        'email': get_cookie('cosmos_email', email),
        'phone': get_cookie('cosmos_phone', ''),
        'info': get_cookie('cosmos_info', ''),
    }

    # Загружаем дополнительные позывные из куки
    other_callsigns_cookie = get_cookie('cosmos_other_callsigns', '')
    if other_callsigns_cookie:
        try:
            saved_callsigns = json.loads(other_callsigns_cookie)
            if isinstance(saved_callsigns, list):
                other_callsigns_list = saved_callsigns
        except json.JSONDecodeError:
            pass

    if request.method == 'POST':
        try:
            # Получаем данные из формы
            main_callsign = request.POST.get('main_callsign', '').strip().upper()
            full_name = request.POST.get('full_name', '').strip()
            email = request.POST.get('email', '').strip()
            phone = request.POST.get('phone', '').strip()
            info = request.POST.get('info', '').strip()

            # Получаем дополнительные позывные (поддерживаем оба формата)
            # Новый формат: массив из input полей
            other_callsigns_from_inputs = request.POST.getlist('other_callsigns_names[]')
            # Старый формат: JSON из hidden field
            other_callsigns_raw = request.POST.get('other_callsigns_json', '[]')

            # Используем данные из input полей если есть, иначе из JSON
            if other_callsigns_from_inputs:
                other_callsigns = [c.strip().upper() for c in other_callsigns_from_inputs if c.strip()]
            else:
                try:
                    other_callsigns = json.loads(other_callsigns_raw)
                    if isinstance(other_callsigns, list):
                        other_callsigns = [c.upper() if isinstance(c, str) else c for c in other_callsigns]
                except (json.JSONDecodeError, TypeError):
                    other_callsigns = []

            # Валидация
            if not main_callsign:
                messages.error(request, 'Позывной обязателен')
                return render(request, 'cosmos_diploma.html', {
                    'form_data': form_data,
                    'other_callsigns': other_callsigns_list,
                })

            if not full_name:
                messages.error(request, 'ФИО обязательно')
                return render(request, 'cosmos_diploma.html', {
                    'form_data': form_data,
                    'other_callsigns': other_callsigns_list,
                })

            if not email:
                messages.error(request, 'Email обязателен')
                return render(request, 'cosmos_diploma.html', {
                    'form_data': form_data,
                    'other_callsigns': other_callsigns_list,
                })

            # Формируем список всех позывных для поиска QSO
            all_callsigns = [main_callsign] + [c for c in other_callsigns if c and c != main_callsign]

            # Запрос QSO из базы данных - ищем по всем позывным
            qso_data = []
            try:
                # Создаем условие WHERE с учетом всех позывных и user_id
                if all_callsigns:
                    placeholders = ','.join(['%s'] * len(all_callsigns))
                    query = f"""
                        SELECT * FROM (
                            SELECT DISTINCT ON (callsign)
                                LEFT(COALESCE(gridsquare, ''), 4) as gridsquare,
                                date, band, time, mode, callsign,
                                rst_sent, rst_rcvd
                            FROM tlog_qso
                            WHERE user_id = %s AND my_callsign IN ({placeholders}) and (prop_mode = 'SAT' OR band = '13CM')
                            ORDER BY callsign, date, time
                        ) AS distinct_qsos
                        ORDER BY date, time
                    """
                    with connection.cursor() as cursor:
                        # Параметры: user_id + все позывные
                        cursor.execute(query, [user.id] + all_callsigns)
                        qso_data = cursor.fetchall()
            except Exception as e:
                messages.warning(request, f'Не удалось получить данные QSO: {str(e)}')
                qso_data = []

            # Проверяем минимальное количество записей
            if len(qso_data) < 100:
                pass  # Предупреждение убрано - сообщение будет показано на странице загрузки файла

            # Путь к шаблону
            template_path = os.path.join(settings.BASE_DIR, 'tlog', 'cosmos', 'Cosmos.docx')

            if not os.path.exists(template_path):
                messages.error(request, 'Шаблон документа не найден')
                return render(request, 'cosmos_diploma.html', {
                    'form_data': form_data,
                    'other_callsigns': other_callsigns_list,
                })

            # Открываем документ
            try:
                from docx import Document
                from docx.shared import Pt
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                
                doc = Document(template_path)
            except ImportError:
                messages.error(request, 'Библиотека python-docx не установлена')
                return render(request, 'cosmos_diploma.html', {
                    'form_data': form_data,
                    'other_callsigns': other_callsigns_list,
                })

            # Форматируем данные
            def format_date(date_val):
                if date_val is None:
                    return ""
                if hasattr(date_val, 'strftime'):
                    return date_val.strftime("%d.%m.%Y")
                return str(date_val)

            def format_time(time_val):
                if time_val is None:
                    return ""
                if hasattr(time_val, 'strftime'):
                    return time_val.strftime("%H:%M")
                return str(time_val)

            def set_cell_border(cell):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                borders = parse_xml(
                    "<w:tcBorders {}><w:top w:val=\"single\" w:sz=\"2\"/><w:left w:val=\"single\" w:sz=\"2\"/><w:right w:val=\"single\" w:sz=\"2\"/><w:bottom w:val=\"single\" w:sz=\"2\"/><w:insideH w:val=\"single\" w:sz=\"2\"/><w:insideV w:val=\"single\" w:sz=\"2\"/></w:tcBorders>".format(
                        nsdecls("w")))
                tcPr.append(borders)

            def set_cell_shading(cell, color):
                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
                cell._tc.get_or_add_tcPr().append(shading_elm)

            def set_col_width(table, col_idx, width):
                for row in table.rows:
                    row.cells[col_idx].width = width

            # Формируем строку с другими позывными
            other_callsigns_str = ', '.join(other_callsigns) if other_callsigns else ''

            # Подсчеты
            count25 = (len(qso_data) // 25) * 25
            current_date = datetime.now().strftime("%d.%m.%Y")
            first_word_name = full_name.split()[0] if full_name else ""

            # Заполняем документ
            for paragraph in doc.paragraphs:
                if "{{YOUR_NAME}}" in paragraph.text:
                    parts = paragraph.text.split("{{YOUR_NAME}}")
                    paragraph.clear()
                    for part in parts[:-1]:
                        paragraph.add_run(part)
                        run = paragraph.add_run(full_name)
                        run.bold = True
                        run.font.size = Pt(14)
                    paragraph.add_run(parts[-1])
                elif "{{MY_CALLSIGN}}" in paragraph.text:
                    parts = paragraph.text.split("{{MY_CALLSIGN}}")
                    paragraph.clear()
                    for part in parts[:-1]:
                        paragraph.add_run(part)
                        run = paragraph.add_run(main_callsign)
                        run.bold = True
                        run.font.size = Pt(14)
                    paragraph.add_run(parts[-1])
                elif "{{OTHER_CALLSIGNS}}" in paragraph.text:
                    parts = paragraph.text.split("{{OTHER_CALLSIGNS}}")
                    paragraph.clear()
                    for part in parts[:-1]:
                        paragraph.add_run(part)
                        run = paragraph.add_run(other_callsigns_str)
                        run.bold = True
                        run.font.size = Pt(14)
                    paragraph.add_run(parts[-1])
                elif "{{YOUR_EMAIL}}" in paragraph.text:
                    parts = paragraph.text.split("{{YOUR_EMAIL}}")
                    paragraph.clear()
                    for part in parts[:-1]:
                        paragraph.add_run(part)
                        run = paragraph.add_run(email)
                        run.bold = True
                        run.font.size = Pt(14)
                    paragraph.add_run(parts[-1])
                elif "{{YOUR_PHONE}}" in paragraph.text:
                    parts = paragraph.text.split("{{YOUR_PHONE}}")
                    paragraph.clear()
                    for part in parts[:-1]:
                        paragraph.add_run(part)
                        run = paragraph.add_run(phone)
                        run.bold = True
                        run.font.size = Pt(14)
                    paragraph.add_run(parts[-1])
                elif "{{INITIAL}}" in paragraph.text:
                    parts = paragraph.text.split("{{INITIAL}}")
                    paragraph.clear()
                    for part in parts[:-1]:
                        paragraph.add_run(part)
                        run = paragraph.add_run(str(len(qso_data)))
                        run.bold = True
                        run.font.size = Pt(14)
                    paragraph.add_run(parts[-1])
                elif "{{INFO}}" in paragraph.text:
                    parts = paragraph.text.split("{{INFO}}")
                    paragraph.clear()
                    for part in parts[:-1]:
                        paragraph.add_run(part)
                        run = paragraph.add_run(info)
                        run.bold = True
                        run.font.size = Pt(14)
                    paragraph.add_run(parts[-1])
                elif "{{BOTTOM_LINE}}" in paragraph.text:
                    paragraph.clear()
                    run = paragraph.add_run("Подпись ")
                    run.font.size = Pt(14)
                    run = paragraph.add_run(first_word_name)
                    run.bold = True
                    run.font.size = Pt(14)
                    run = paragraph.add_run("       Позывной ")
                    run.font.size = Pt(14)
                    run = paragraph.add_run(main_callsign)
                    run.bold = True
                    run.font.size = Pt(14)
                    run = paragraph.add_run("          Дата ")
                    run.font.size = Pt(14)
                    run = paragraph.add_run(current_date)
                    run.bold = True
                    run.font.size = Pt(14)
                elif "{{count25}}" in paragraph.text:
                    original_run = paragraph.runs[0] if paragraph.runs else None
                    original_font_size = original_run.font.size if original_run else None
                    original_font_name = original_run.font.name if original_run else None
                    original_bold = original_run.bold if original_run else None

                    parts = paragraph.text.split("{{count25}}")
                    paragraph.clear()
                    for part in parts[:-1]:
                        paragraph.add_run(part)
                        if original_font_size:
                            paragraph.runs[-1].font.size = original_font_size
                        if original_font_name:
                            paragraph.runs[-1].font.name = original_font_name
                        if original_bold:
                            paragraph.runs[-1].bold = True
                    run = paragraph.add_run(str(count25))
                    run.bold = True
                    run.font.size = Pt(22)
                    paragraph.add_run(parts[-1])
                    if original_font_size:
                        paragraph.runs[-1].font.size = original_font_size
                    if original_font_name:
                        paragraph.runs[-1].font.name = original_font_name
                    if original_bold:
                        paragraph.runs[-1].bold = True
                elif "{{TABLE}}" in paragraph.text:
                    paragraph.clear()
                    if qso_data:
                        table = doc.add_table(rows=1, cols=9)
                        headers = ["№", "QTH-loc", "Дата", "Диапазон", "Время", "Модуляция", "Позывной", "Рапорт переданный", "Рапорт принятый"]
                        header_cells = table.rows[0].cells
                        for i, header in enumerate(headers):
                            header_cells[i].text = header
                            run = header_cells[i].paragraphs[0].runs[0]
                            run.bold = True
                            run.font.size = Pt(10)
                            set_cell_border(header_cells[i])
                            set_cell_shading(header_cells[i], "D9D9D9")
                        for row_num, row_data in enumerate(qso_data, 1):
                            row = table.add_row().cells
                            row[0].text = str(row_num)
                            row[1].text = row_data[0] or ""
                            row[2].text = format_date(row_data[1])
                            row[3].text = row_data[2] or ""
                            row[4].text = format_time(row_data[3])
                            row[5].text = row_data[4] or ""
                            row[6].text = row_data[5] or ""
                            row[7].text = row_data[6] or ""
                            row[8].text = row_data[7] or ""
                            for cell in row:
                                cell.paragraphs[0].runs[0].font.size = Pt(10)
                                set_cell_border(cell)
                        set_col_width(table, 0, Pt(10))

            # Формируем имя файла
            safe_username = user.username.replace(' ', '_')
            output_filename = f"Заявка_Cosmos_{safe_username}.docx"

            # Сохраняем во временный буфер
            from io import BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Сохраняем файл во временный файл для скачивания
            temp_dir = os.path.join(settings.BASE_DIR, 'temp')
            os.makedirs(temp_dir, exist_ok=True)
            temp_file_path = os.path.join(temp_dir, f'cosmos_{user.id}_{int(time.time())}.docx')

            with open(temp_file_path, 'wb') as f:
                f.write(buffer.getvalue())

            # Сохраняем путь к файлу в сессии
            request.session['cosmos_download_path'] = temp_file_path
            request.session['cosmos_download_filename'] = output_filename

            # Сохраняем данные формы в сессии
            request.session['cosmos_form_data'] = {
                'main_callsign': main_callsign,
                'full_name': full_name,
                'email': email,
                'phone': phone,
                'info': info,
                'other_callsigns': other_callsigns,
            }

            # Формируем сообщение
            if len(qso_data) >= 100:
                messages.success(request, f'✅ Заявка сформирована! Найдено {len(qso_data)} уникальных QSO.')
            else:
                messages.warning(request, f'⚠️ Внимание! В заявке найдено только {len(qso_data)} уникальных QSO. Для диплома необходимо минимум 100.')

            # Редирект на страницу с сообщением
            return redirect('cosmos_diploma')

        except Exception as e:
            messages.error(request, f'Ошибка при формировании заявки: {str(e)}')
            return render(request, 'cosmos_diploma.html', {
                'form_data': form_data,
                'other_callsigns': other_callsigns_list,
            })

    # GET запрос - показываем форму
    # Проверяем, есть ли готовый файл для скачивания
    has_cosmos_file = 'cosmos_download_path' in request.session

    # Загружаем данные формы из сессии если есть
    form_data_session = request.session.get('cosmos_form_data')
    if form_data_session:
        form_data['main_callsign'] = form_data_session.get('main_callsign', main_callsign)
        form_data['full_name'] = form_data_session.get('full_name', full_name)
        form_data['email'] = form_data_session.get('email', email)
        form_data['phone'] = form_data_session.get('phone', '')
        form_data['info'] = form_data_session.get('info', '')
        # Загружаем дополнительные позывные из сессии
        other_callsigns_from_session = form_data_session.get('other_callsigns', [])
        if other_callsigns_from_session:
            other_callsigns_list = other_callsigns_from_session

    return render(request, 'cosmos_diploma.html', {
        'form_data': form_data,
        'other_callsigns': other_callsigns_list,
        'page_title': 'Заявка на диплом Cosmos',
        'page_subtitle': 'Формирование заявки на диплом Cosmos (СРР)',
        'has_cosmos_file': has_cosmos_file,
    })


@login_required
def cosmos_download(request):
    """
    Скачивание сформированного файла заявки Cosmos
    """
    if not request.user.is_authenticated:
        return redirect('login_page')

    download_path = request.session.get('cosmos_download_path')
    download_filename = request.session.get('cosmos_download_filename')

    if not download_path or not os.path.exists(download_path):
        messages.error(request, 'Файл заявки не найден. Пожалуйста, сформируйте заявку заново.')
        return redirect('cosmos_diploma')

    try:
        with open(download_path, 'rb') as f:
            content = f.read()

        # Удаляем временный файл
        os.remove(download_path)

        # Очищаем сессию
        del request.session['cosmos_download_path']
        del request.session['cosmos_download_filename']

        response = HttpResponse(
            content,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        from django.utils.encoding import escape_uri_path
        response['Content-Disposition'] = f'attachment; filename="{escape_uri_path(download_filename)}"'
        return response

    except Exception as e:
        messages.error(request, f'Ошибка при скачивании файла: {str(e)}')
        return redirect('cosmos_diploma')