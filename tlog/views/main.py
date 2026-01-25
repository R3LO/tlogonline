

# Main views (home, dashboard, profile)


from django.shortcuts import render, redirect
from django.contrib import messages
from django.db.models import Q
from django.http import JsonResponse, HttpResponse, HttpResponseRedirect
from django.urls import reverse
from django.contrib.auth.models import User
from django.utils import timezone
from datetime import timedelta
import json
import time
import os
from ..models import QSO, RadioProfile, ADIFUpload, check_user_blocked, ChatMessage


def home(request):
    """
    Главная страница
    """
    total_users = User.objects.count()
    total_qso = QSO.objects.count()

    return render(request, 'index.html', {
        'total_users': total_users,
        'total_qso': total_qso,
    })


def get_callsigns_list(request):
    """
    Получение списка уникальных позывных пользователей для автодополнения
    """
    query = request.GET.get('q', '').upper()
    if len(query) < 1:
        return JsonResponse({'callsigns': []})

    callsigns_set = set()

    # 1. Ищем в поле callsign модели RadioProfile
    profiles_with_callsign = RadioProfile.objects.filter(
        callsign__icontains=query
    ).exclude(callsign='').values_list('callsign', flat=True).distinct()
    callsigns_set.update([c.upper() for c in profiles_with_callsign])

    # 2. Ищем в поле my_callsigns (JSON) модели RadioProfile
    profiles = RadioProfile.objects.filter(
        my_callsigns__isnull=False
    ).exclude(
        my_callsigns=[]
    ).exclude(
        my_callsigns=''
    )

    import json
    for profile in profiles:
        try:
            my_callsigns = json.loads(profile.my_callsigns) if isinstance(profile.my_callsigns, str) else profile.my_callsigns
            if isinstance(my_callsigns, list):
                for item in my_callsigns:
                    if isinstance(item, dict) and 'name' in item:
                        name = item['name'].upper()
                        if query in name:
                            callsigns_set.add(name)
                    elif isinstance(item, str):
                        name = item.upper()
                        if query in name:
                            callsigns_set.add(name)
        except (json.JSONDecodeError, TypeError):
            pass

    # 3. Также ищем в QSO (my_callsign)
    qso_callsigns = QSO.objects.filter(
        my_callsign__icontains=query
    ).values_list('my_callsign', flat=True).distinct()
    callsigns_set.update([c.upper() for c in qso_callsigns])

    # Ограничиваем до 10 результатов и сортируем
    callsigns_list = sorted(list(callsigns_set))[:10]

    return JsonResponse({'callsigns': callsigns_list})


def dashboard(request):
    """
    Личный кабинет радиолюбителя
    """
    if not request.user.is_authenticated:
        return redirect('login_page')

    # Проверяем, не заблокирован ли пользователь
    is_blocked, reason = check_user_blocked(request.user)
    if is_blocked:
        return render(request, 'blocked.html', {'reason': reason})

    # Обработка формы ручного ввода QSO
    if request.method == 'POST' and request.POST.get('action') == 'add_qso':
        try:
            from datetime import date, time
            from django.db.utils import IntegrityError

            # Получаем данные из формы
            my_callsign = request.POST.get('my_callsign', '').strip().upper()
            callsign = request.POST.get('callsign', '').strip().upper()
            date_str = request.POST.get('date', '').strip()
            time_str = request.POST.get('time', '').strip()
            band = request.POST.get('band', '').strip().upper()
            mode = request.POST.get('mode', '').strip().upper()
            rst_rcvd = request.POST.get('rst_rcvd', '').strip().upper()
            rst_sent = request.POST.get('rst_sent', '').strip().upper()
            gridsquare = request.POST.get('gridsquare', '').strip().upper()
            ru_region = request.POST.get('ru_region', '').strip().upper()

            # Валидация обязательных полей
            if not all([my_callsign, callsign, date_str, time_str, band, mode]):
                messages.error(request, 'Все обязательные поля должны быть заполнены')
                return redirect('dashboard')

            # Валидация форматов
            if len(my_callsign) > 20 or len(callsign) > 20:
                messages.error(request, 'Позывной не должен превышать 20 символов')
                return redirect('dashboard')

            if len(gridsquare) > 8:
                messages.error(request, 'QTH локатор не должен превышать 8 символов')
                return redirect('dashboard')

            # Парсинг даты и времени
            try:
                qso_date = date.fromisoformat(date_str)
                qso_time = time.fromisoformat(time_str)
            except ValueError:
                messages.error(request, 'Неверный формат даты или времени')
                return redirect('dashboard')

            # Проверка на дубликат
            duplicate_exists = QSO.objects.filter(
                user=request.user,
                my_callsign=my_callsign,
                callsign=callsign,
                date=qso_date,
                time=qso_time,
                band=band,
                mode=mode
            ).exists()

            if duplicate_exists:
                messages.warning(request, 'QSO с такими данными уже существует в логе')
                return redirect('dashboard')

            # Определение частоты по диапазону (упрощенное сопоставление)
            band_frequencies = {
                '160m': 1.9, '80m': 3.7, '40m': 7.1, '30m': 10.15, '20m': 14.2,
                '17m': 18.12, '15m': 21.3, '12m': 24.95, '10m': 28.5, '6m': 52.0,
                '2m': 144.0, '70cm': 435.0, '23cm': 1290.0, '13cm': 2400.0, '3cm': 10000.0
            }

            frequency = band_frequencies.get(band, 0.0)

            # Пересчитываем cqz, ituz, continent, r150s, dxcc, ru_region по позывному
            cqz = None
            ituz = None
            continent = None
            r150s_country = None
            dxcc = None
            ru_region_val = None

            if callsign:
                import os
                from django.conf import settings
                from tlog import r150s
                from tlog.region_ru import RussianRegionFinder

                db_path = os.path.join(settings.BASE_DIR, 'tlog', 'r150cty.dat')
                cty_path = os.path.join(settings.BASE_DIR, 'tlog', 'cty.dat')

                r150s.init_database(db_path)
                r150s.init_cty_database(cty_path)

                dxcc_info = r150s.get_dxcc_info(callsign, db_path)
                if dxcc_info:
                    cqz = dxcc_info.get('cq_zone')
                    ituz = dxcc_info.get('itu_zone')
                    continent = dxcc_info.get('continent')

                    r150s_country = dxcc_info.get('country')

                    dxcc = r150s.get_cty_primary_prefix(callsign, cty_path)

                # Определяем код региона России только для российских позывных (UA, UA9, UA2)
                if dxcc and dxcc.upper() in ('UA', 'UA9', 'UA2'):
                    exceptions_path = os.path.join(settings.BASE_DIR, 'tlog', 'exceptions.dat')
                    region_finder = RussianRegionFinder(exceptions_file=exceptions_path)
                    ru_region_val = region_finder.get_region_code(callsign)

            # Создание записи QSO
            try:
                QSO.objects.create(
                    user=request.user,
                    my_callsign=my_callsign,
                    callsign=callsign,
                    date=qso_date,
                    time=qso_time,
                    frequency=frequency,
                    band=band,
                    mode=mode,
                    rst_rcvd=rst_rcvd if rst_rcvd else None,
                    rst_sent=rst_sent if rst_sent else None,
                    gridsquare=gridsquare if gridsquare else None,
                    ru_region=ru_region_val,
                    cqz=cqz,
                    ituz=ituz,
                    continent=continent if continent else None,
                    r150s=r150s_country if r150s_country else None,
                    dxcc=dxcc if dxcc else None,
                    lotw='N',
                    paper_qsl='N'
                )

                messages.success(request, f'QSO с {callsign} успешно добавлено в лог')
                return redirect('dashboard')

            except IntegrityError as e:
                messages.error(request, f'Ошибка при сохранении QSO: {str(e)}')
                return redirect('dashboard')

        except Exception as e:
            messages.error(request, f'Ошибка при обработке формы: {str(e)}')
            return redirect('dashboard')

    # Получаем профиль радиолюбителя
    try:
        profile = request.user.radio_profile
    except RadioProfile.DoesNotExist:
        profile = None

    # Получаем QSO пользователя
    user_qso = QSO.objects.filter(user=request.user)

    # Подсчитываем статистику
    total_qso = user_qso.count()
    unique_callsigns = user_qso.values('callsign').distinct().count()

    # Статистика по DXCC (страны)
    dxcc_count = user_qso.exclude(dxcc__isnull=True).exclude(dxcc='').values('dxcc').distinct().count()

    # Статистика по Р-150-С (регионы России)
    r150s_count = user_qso.exclude(r150s__isnull=True).exclude(r150s='').values('r150s').distinct().count()

    # Статистика по регионам России (ru_region)
    ru_region_count = user_qso.exclude(ru_region__isnull=True).exclude(ru_region='').values('ru_region').distinct().count()

    # Статистика по видам модуляции
    mode_stats = {}
    mode_choices = dict(QSO._meta.get_field('mode').choices)
    for mode in mode_choices.keys():
        count = user_qso.filter(mode=mode).count()
        if count > 0:
            mode_stats[mode] = count

    # Последние QSO (последние 10)
    recent_qso = user_qso.order_by('-date', '-time')[:10]

    # Получаем загруженные ADIF файлы
    adif_uploads = ADIFUpload.objects.filter(user=request.user).order_by('-upload_date')[:10]

    context = {
        'user': request.user,
        'profile': profile,
        'total_qso': total_qso,
        'unique_callsigns': unique_callsigns,
        'dxcc_count': dxcc_count,
        'r150s_count': r150s_count,
        'ru_region_count': ru_region_count,
        'recent_qso': recent_qso,
        'mode_statistics': mode_stats,
        'adif_uploads': adif_uploads,
    }

    return render(request, 'dashboard.html', context)


def profile_update(request):
    """
    Обновление профиля радиолюбителя (Django 5.2)
    """
    if not request.user.is_authenticated:
        return redirect('login_page')

    # Проверяем, не заблокирован ли пользователь
    is_blocked, reason = check_user_blocked(request.user)
    if is_blocked:
        return render(request, 'blocked.html', {'reason': reason})

    # Получаем или создаем профиль
    try:
        profile = RadioProfile.objects.get(user=request.user)
    except RadioProfile.DoesNotExist:
        profile = RadioProfile.objects.create(user=request.user)

    if request.method == 'POST':
        try:
            # Обновляем поля профиля (callsign всегда равен username)
            profile.callsign = request.user.username.upper()
            profile.first_name = request.POST.get('first_name', '').strip()
            profile.last_name = request.POST.get('last_name', '').strip()
            profile.qth = request.POST.get('qth', '').strip()
            profile.my_gridsquare = request.POST.get('my_gridsquare', '').strip().upper()

            # Обновляем email пользователя
            new_email = request.POST.get('email', '').strip()
            if new_email:
                # Простая валидация email
                import re
                email_pattern = r'^[^\s@]+@[^\s@]+\.[^\s@]+$'
                if re.match(email_pattern, new_email):
                    request.user.email = new_email
                    request.user.save(update_fields=['email'])
                else:
                    messages.error(request, 'Введите корректный email адрес')
                    return render(request, 'profile_edit.html', {
                        'profile': profile,
                    })

            # Обработка настроек LoTW
            use_lotw = 'use_lotw' in request.POST
            if use_lotw:
                profile.lotw_user = request.POST.get('lotw_user', '').strip()
                profile.lotw_password = request.POST.get('lotw_password', '').strip()
                # lotw_chk_pass сохраняется как есть (обновляется при проверке)
            else:
                # Очищаем данные LoTW если чекбокс не выбран
                profile.lotw_user = ''
                profile.lotw_password = ''
                profile.lotw_chk_pass = False

            # Обрабатываем my_callsigns из JSON
            my_callsigns_json = request.POST.get('my_callsigns_json', '[]')
            try:
                new_my_callsigns = json.loads(my_callsigns_json)
            except json.JSONDecodeError:
                new_my_callsigns = []

            # Проверяем, изменились ли my_callsigns
            old_my_callsigns = profile.my_callsigns if profile.my_callsigns else []
            if isinstance(old_my_callsigns, str):
                old_my_callsigns = json.loads(old_my_callsigns)

            # Если изменились - очищаем lotw_lastsync
            if old_my_callsigns != new_my_callsigns:
                profile.lotw_lastsync = None
                profile.my_callsigns = new_my_callsigns
                profile.save(update_fields=['lotw_lastsync', 'my_callsigns'])
            else:
                profile.my_callsigns = new_my_callsigns
                profile.save()

            # Также обновляем User модель
            request.user.first_name = profile.first_name
            request.user.last_name = profile.last_name
            request.user.save(update_fields=['first_name', 'last_name'])

            messages.success(request, 'Профиль успешно обновлён')
            return redirect('profile_update')
        except Exception as e:
            messages.error(request, f'Ошибка при обновлении профиля: {str(e)}')

    # Для GET запроса или после POST с ошибкой - показываем форму
    return render(request, 'profile_edit.html', {
        'profile': profile,
    })


def verify_lotw_credentials(request):
    """
    Проверка логина и пароля LoTW
    """
    if not request.user.is_authenticated:
        return JsonResponse({'error': 'Вы должны быть авторизованы'}, status=403)

    if request.method != 'POST':
        return JsonResponse({'error': 'Метод не поддерживается'}, status=405)

    try:
        import requests

        # Получаем логин и пароль из POST данных
        data = json.loads(request.body)
        login = data.get('login', '').strip()
        password = data.get('password', '').strip()

        if not login or not password:
            return JsonResponse({'error': 'Логин и пароль обязательны'}, status=400)

        # Функция проверки
        def check_lotw_pass(login, password):
            params = {
                'login': login,
                'password': password,
            }
            response = requests.get(
                "https://lotw.arrl.org/lotwuser/lotwreport.adi",
                params=params,
                timeout=15
            )
            if response.text.strip().startswith('ARRL Logbook of the World Status Report'):
                return True
            elif '<HTML>' in response.text.upper() or '<!DOCTYPE HTML' in response.text.upper():
                return False
            return False

        # Выполняем проверку
        is_valid = check_lotw_pass(login, password)

        # Обновляем профиль пользователя
        try:
            profile = RadioProfile.objects.get(user=request.user)
            profile.lotw_chk_pass = is_valid
            if is_valid:
                profile.lotw_user = login
                profile.lotw_password = password
            profile.save()
        except RadioProfile.DoesNotExist:
            pass

        return JsonResponse({
            'success': True,
            'is_valid': is_valid,
            'message': 'Логин и пароль верны' if is_valid else 'Логин или пароль неверны'
        })

    except json.JSONDecodeError:
        return JsonResponse({'error': 'Неверный формат данных'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


def delete_lotw_credentials(request):
    """
    Удаление логина и пароля LoTW
    """
    if not request.user.is_authenticated:
        return JsonResponse({'error': 'Вы должны быть авторизованы'}, status=403)

    if request.method != 'POST':
        return JsonResponse({'error': 'Метод не поддерживается'}, status=405)

    try:
        profile = RadioProfile.objects.get(user=request.user)
        profile.lotw_user = ''
        profile.lotw_password = ''
        profile.lotw_chk_pass = False
        profile.save()

        return JsonResponse({'success': True})

    except RadioProfile.DoesNotExist:
        return JsonResponse({'error': 'Профиль не найден'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


def change_password(request):
    """
    Смена пароля пользователя
    """
    if not request.user.is_authenticated:
        return JsonResponse({'error': 'Вы должны быть авторизованы'}, status=403)

    if request.method != 'POST':
        return JsonResponse({'error': 'Метод не поддерживается'}, status=405)

    try:
        data = json.loads(request.body)
        new_password = data.get('password', '')

        if not new_password:
            return JsonResponse({'error': 'Пароль не может быть пустым'}, status=400)

        if len(new_password) < 8:
            return JsonResponse({'error': 'Пароль должен содержать минимум 8 символов'}, status=400)

        # Устанавливаем новый пароль через set_password (Django автоматически хеширует)
        request.user.set_password(new_password)
        request.user.save()

        # Разлогиниваем пользователя и перенаправляем на страницу логина
        from django.contrib.auth import logout
        logout(request)

        return JsonResponse({
            'success': True,
            'message': 'Пароль успешно изменён',
            'redirect': '/login/'
        })

    except json.JSONDecodeError:
        return JsonResponse({'error': 'Неверный формат данных'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


def chat_list(request):
    """
    Получение списка последних сообщений чата (для AJAX)
    """
    # Получаем последние 100 сообщений
    messages = ChatMessage.objects.all()[:100]

    messages_data = []
    for msg in messages:
        messages_data.append({
            'id': str(msg.id),
            'user_id': msg.user.id,
            'username': msg.username,
            'message': msg.message,
            'created_at': msg.created_at.strftime('%H:%M'),
        })

    return JsonResponse({'messages': messages_data})


def chat_send(request):
    """
    Отправка нового сообщения в чат
    """
    if not request.user.is_authenticated:
        return JsonResponse({'error': 'Вы должны быть авторизованы'}, status=403)

    if request.method != 'POST':
        return JsonResponse({'error': 'Метод не поддерживается'}, status=405)

    try:
        import json
        data = json.loads(request.body)
        message_text = data.get('message', '').strip()

        if not message_text:
            return JsonResponse({'error': 'Сообщение не может быть пустым'}, status=400)

        if len(message_text) > 500:
            return JsonResponse({'error': 'Сообщение слишком длинное (максимум 500 символов)'}, status=400)

        # Получаем callsign пользователя
        try:
            profile = request.user.radio_profile
            username = profile.callsign or request.user.username
        except RadioProfile.DoesNotExist:
            username = request.user.username

        # Создаем сообщение
        chat_message = ChatMessage.objects.create(
            user=request.user,
            username=username,
            message=message_text
        )

        # Удаляем старые сообщения, оставляем только последние 100
        # Получаем ID последних 100 сообщений
        keep_ids = list(ChatMessage.objects.order_by('-created_at')[:100].values_list('id', flat=True))
        # Удаляем все остальные
        ChatMessage.objects.exclude(id__in=keep_ids).delete()

        return JsonResponse({
            'success': True,
            'message': {
                'id': str(chat_message.id),
                'username': chat_message.username,
                'message': chat_message.message,
                'created_at': chat_message.created_at.strftime('%H:%M'),
            }
        })

    except json.JSONDecodeError:
        return JsonResponse({'error': 'Неверный формат данных'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


def custom_404_view(request, exception):
    """
    Обработчик ошибки 404 - страница не найдена
    """
    return render(request, '404.html', {
        'request_path': request.path,
    }, status=404)


def qo100_regions(request):
    """
    Страница рейтинга QO-100 - регионы России
    Показывает статистику по регионам РФ для всех QSO с lotw = 'Y'
    """
    from tlog.region_ru import RussianRegionFinder

    region_finder = RussianRegionFinder()

    # Получаем уникальные пары my_callsign + регион + callsign для QSO с lotw = 'Y'
    qso_filtered = QSO.objects.filter(
        lotw='Y',
        ru_region__isnull=False
    ).exclude(ru_region='').values('my_callsign', 'ru_region', 'callsign').distinct()

    # Группируем по my_callsign, затем по региону
    from collections import defaultdict
    callsign_data = defaultdict(lambda: defaultdict(set))

    for item in qso_filtered:
        my_call = item['my_callsign']
        region_code = item['ru_region']
        call = item['callsign']
        callsign_data[my_call][region_code].add(call)

    # Формируем список с позывным, количеством и данными регионов
    ratings = []
    for my_call, regions_dict in callsign_data.items():
        regions_list = []
        for region_code, callsigns in regions_dict.items():
            region_name = region_finder.region_data.get(region_code, region_code)
            regions_list.append({
                'code': region_code,
                'name': region_name,
                'callsigns': sorted(list(callsigns))
            })
        # Сортируем по названию региона
        regions_list.sort(key=lambda x: x['name'])

        ratings.append({
            'callsign': my_call,
            'count': len(regions_list),
            'regions': regions_list
        })

    # Сортируем по количеству (убывание), затем по позывному
    ratings.sort(key=lambda x: (-x['count'], x['callsign']))

    return render(request, 'qo100/regions.html', {
        'ratings': ratings,
        'page_title': 'Регионы России QO-100',
        'page_subtitle': 'По регионам РФ, подтвержденным в LoTW',
    })


def qo100_r150s(request):
    """
    Страница рейтинга QO-100 - страны Р-150-С
    Показывает статистику по странам для всех QSO с lotw = 'Y'
    """
    from tlog.r150s import DXCCDatabase
    import os

    # Загружаем базу данных R150
    db_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'r150cty.dat')
    dxcc_db = DXCCDatabase(db_path)

    # Создаем словарь код -> название страны
    country_names = {}
    for entry in dxcc_db.entries:
        # Проверяем все префиксы, чтобы найти нужный код
        for prefix in entry.prefixes:
            # Код r150s может быть в формате "=P" или просто префикс
            if '=' in prefix:
                code = prefix.split('=')[1].strip()
                country_names[code] = entry.name

    # Получаем уникальные пары my_callsign + r150s для QSO с lotw = 'Y'
    qso_filtered = QSO.objects.filter(
        lotw='Y',
        r150s__isnull=False
    ).exclude(r150s='').values('my_callsign', 'r150s', 'callsign').distinct()

    # Группируем по my_callsign, затем по стране
    from collections import defaultdict
    callsign_data = defaultdict(lambda: defaultdict(set))

    for item in qso_filtered:
        my_call = item['my_callsign']
        country_code = item['r150s']
        call = item['callsign']
        callsign_data[my_call][country_code].add(call)

    # Формируем список с позывным, количеством и данными стран
    ratings = []
    for my_call, countries_dict in callsign_data.items():
        countries_list = []
        for country_code, callsigns in countries_dict.items():
            country_name = country_names.get(country_code, country_code)
            countries_list.append({
                'code': country_code,
                'name': country_name,
                'callsigns': sorted(list(callsigns))
            })
        # Сортируем по названию страны
        countries_list.sort(key=lambda x: x['name'])

        ratings.append({
            'callsign': my_call,
            'count': len(countries_list),
            'countries': countries_list
        })

    # Сортируем по количеству (убывание), затем по позывному
    ratings.sort(key=lambda x: (-x['count'], x['callsign']))

    return render(request, 'qo100/r150s.html', {
        'ratings': ratings,
        'page_title': 'Страны Р-150-С QO-100',
        'page_subtitle': 'По странам Р-150-С, подтвержденным в LoTW',
    })


def qo100_grids(request):
    """
    Страница рейтинга QO-100 - локаторы
    Показывает статистику по локаторам для всех QSO с lotw = 'Y'
    """
    # Получаем уникальные пары my_callsign + gridsquare для QSO с lotw = 'Y'
    qso_filtered = QSO.objects.filter(
        lotw='Y',
        gridsquare__isnull=False
    ).exclude(gridsquare='').values('my_callsign', 'gridsquare', 'callsign').distinct()

    # Группируем по my_callsign, затем по локатору (первые 4 знака)
    from collections import defaultdict
    callsign_data = defaultdict(lambda: defaultdict(set))

    for item in qso_filtered:
        my_call = item['my_callsign']
        # Берем только первые 4 знака локатора
        grid_full = item['gridsquare']
        grid_short = grid_full[:4] if len(grid_full) >= 4 else grid_full
        call = item['callsign']
        callsign_data[my_call][grid_short].add(call)

    # Формируем список с позывным, количеством и данными локаторов
    ratings = []
    for my_call, grids_dict in callsign_data.items():
        grids_list = []
        for grid_code, callsigns in grids_dict.items():
            grids_list.append({
                'code': grid_code,
                'callsigns': sorted(list(callsigns))
            })
        # Сортируем по коду локатора
        grids_list.sort(key=lambda x: x['code'])

        ratings.append({
            'callsign': my_call,
            'count': len(grids_list),
            'grids': grids_list
        })

    # Сортируем по количеству (убывание), затем по позывному
    ratings.sort(key=lambda x: (-x['count'], x['callsign']))

    return render(request, 'qo100/grids.html', {
        'ratings': ratings,
        'page_title': 'Локаторы QO-100',
        'page_subtitle': 'По локаторам, подтвержденным в LoTW',
    })


def qo100_unique_callsigns(request):
    """
    Страница рейтинга QO-100 - уникальные позывные
    Показывает статистику по уникальным позывным для всех QSO с lotw = 'Y'
    """
    # Получаем уникальные пары my_callsign + callsign для QSO с lotw = 'Y'
    qso_filtered = QSO.objects.filter(
        lotw='Y',
        callsign__isnull=False
    ).exclude(callsign='').values('my_callsign', 'callsign').distinct()

    # Группируем по my_callsign и собираем уникальные позывные
    from collections import defaultdict
    callsign_data = defaultdict(list)

    for item in qso_filtered:
        my_call = item['my_callsign']
        call = item['callsign']
        if call not in callsign_data[my_call]:
            callsign_data[my_call].append(call)

    # Формируем список с позывным и списком уникальных корреспондентов
    ratings = []
    for my_call, callsigns in callsign_data.items():
        ratings.append({
            'callsign': my_call,
            'count': len(callsigns),
            'unique_callsigns': sorted(callsigns)  # Сортировка по алфавиту
        })

    # Сортируем по количеству (убывание), затем по позывному
    ratings.sort(key=lambda x: (-x['count'], x['callsign']))

    return render(request, 'qo100/unique_callsigns.html', {
        'ratings': ratings,
        'page_title': 'Уникальные позывные QO-100',
        'page_subtitle': 'По уникальным позывным, подтвержденным в LoTW',
    })


def qo100_dxcc(request):
    """
    Страница рейтинга QO-100 - страны DXCC
    Показывает статистику по странам DXCC для всех QSO с lotw = 'Y'
    """
    from tlog.r150s import CTYDatabase
    import os

    # Загружаем базу данных CTY
    cty_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'cty.dat')
    cty_db = CTYDatabase(cty_path)

    # Создаем словарь код (primary_prefix) -> название страны
    country_names = {}
    for entry in cty_db.entries:
        primary_prefix = entry.primary_prefix
        if primary_prefix:
            country_names[primary_prefix] = entry.name

    # Получаем уникальные пары my_callsign + dxcc + callsign для QSO с lotw = 'Y'
    qso_filtered = QSO.objects.filter(
        lotw='Y',
        dxcc__isnull=False
    ).exclude(dxcc='').values('my_callsign', 'dxcc', 'callsign').distinct()

    # Группируем по my_callsign, затем по стране DXCC
    from collections import defaultdict
    callsign_data = defaultdict(lambda: defaultdict(set))

    for item in qso_filtered:
        my_call = item['my_callsign']
        dxcc_code = item['dxcc']
        call = item['callsign']
        callsign_data[my_call][dxcc_code].add(call)

    # Формируем список с позывным, количеством и данными стран
    ratings = []
    for my_call, countries_dict in callsign_data.items():
        countries_list = []
        for country_code, callsigns in countries_dict.items():
            country_name = country_names.get(country_code, country_code)
            countries_list.append({
                'code': country_code,
                'name': country_name,
                'callsigns': sorted(list(callsigns))
            })
        # Сортируем по названию страны
        countries_list.sort(key=lambda x: x['name'])

        ratings.append({
            'callsign': my_call,
            'count': len(countries_list),
            'countries': countries_list
        })

    # Сортируем по количеству (убывание), затем по позывному
    ratings.sort(key=lambda x: (-x['count'], x['callsign']))

    return render(request, 'qo100/dxcc.html', {
        'ratings': ratings,
        'page_title': 'Страны DXCC QO-100',
        'page_subtitle': 'По странам DXCC, подтвержденным в LoTW',
    })


def qo100_converter(request):
    """
    Конвертер ADIF файлов для QO-100
    """
    if not request.user.is_authenticated:
        return redirect('login_page')

    # Проверяем, не заблокирован ли пользователь
    is_blocked, reason = check_user_blocked(request.user)
    if is_blocked:
        return render(request, 'blocked.html', {'reason': reason})

    preview = None
    download_url = None
    download_filename = None

    if request.method == 'POST':
        try:
            import re
            import uuid
            import os
            from django.conf import settings
            from django.http import HttpResponse

            adif_file = request.FILES.get('adif_file')
            if not adif_file:
                messages.error(request, 'Пожалуйста, выберите файл для загрузки')
                return render(request, 'qo100/converter.html')

            # Читаем содержимое файла
            content = adif_file.read().decode('utf-8', errors='ignore')

            # Переводим в верхний регистр для统一 обработки
            content = content.upper()

            # Проверяем наличие обязательного тега EOH
            if '<EOH>' not in content:
                messages.error(request, 'В файле отсутствует обязательный тег <EOH>')
                return render(request, 'qo100/converter.html')

            # Удаляем все что до <EOH> (включая старую шапку), но <EOH> оставляем
            content = content.split('<EOH>', 1)[1]

            # Шаблон для поиска тегов ADIF
            # Формат: <TAG_NAME:LENGTH>VALUE
            tag_pattern = re.compile(r'<([A-Z_]+):(\d+)>([^<]*)')

            # Проверяем наличие обязательных тегов во всех записях
            records = content.split('<EOR>')
            missing_tags = []
            valid_records = []

            for i, record in enumerate(records):
                record = record.strip()
                if not record:
                    continue

                # Ищем теги в записи
                tags = {}
                for match in tag_pattern.finditer(record):
                    tag_name = match.group(1)
                    tag_length = int(match.group(2))
                    tag_value = match.group(3)
                    tags[tag_name] = tag_value

                # Проверяем обязательные теги: QSO_DATE, TIME_OFF или TIME_ON, MODE
                has_qso_date = 'QSO_DATE' in tags
                has_time = 'TIME_OFF' in tags or 'TIME_ON' in tags
                has_mode = 'MODE' in tags

                if not (has_qso_date and has_time and has_mode):
                    if i not in missing_tags:
                        missing_tags.append(i + 1)
                else:
                    valid_records.append({
                        'index': i + 1,
                        'tags': tags,
                        'original': record
                    })

            if not valid_records:
                messages.error(request, 'В файле не найдено корректных записей QSO')
                return render(request, 'qo100/converter.html')

            if missing_tags:
                messages.warning(request,
                    f'В записях №{", ".join(map(str, missing_tags))} отсутствуют обязательные теги '
                    '(QSO_DATE, TIME_OFF/TIME_ON или MODE). Эти записи будут пропущены.')

            # Конвертируем записи
            converted_records = []
            for record_data in valid_records:
                record = record_data['original']
                tags = record_data['tags']

                # Удаляем тег FREQ если есть (заменяем на пробел)
                record = re.sub(r'\s*<FREQ:\d+>[^<]*', ' ', record)

                # Удаляем старые PROP_MODE и SAT_NAME если есть (заменяем на пробел)
                record = re.sub(r'\s*<PROP_MODE:\d+>[^<]*', ' ', record)
                record = re.sub(r'\s*<SAT_NAME:\d+>[^<]*', ' ', record)

                # Заменяем или добавляем тег BAND на <BAND:4>13CM (всегда один экземпляр, с пробелом)
                record = re.sub(r'\s*<BAND:\d+>[^<]*', ' ', record)
                new_band = '<BAND:4>13CM'

                # Добавляем BAND, PROP_MODE и SAT_NAME перед EOR (с пробелами между тегами)
                # Убираем лишние пробелы и добавляем пробел перед каждым тегом
                record = re.sub(r'\s+', ' ', record).strip()
                record = record + ' ' + new_band + ' <PROP_MODE:3>SAT <SAT_NAME:6>QO-100 <EOR>'

                converted_records.append(record)

            # Формируем итоговый файл с нашей шапкой
            adif_header = "ADIF Converter from TLogOnline.com\nCopyright Vladimir R3LO\n<PROGRAMID:14>TlogOnline.com\n<EOH>\n"
            converted_content = adif_header + '\n'.join(converted_records)

            # Формируем имя файла с timestamp
            from django.utils import timezone
            timestamp = timezone.now().strftime('%Y%m%d_%H%M')
            safe_username = request.user.username.replace(' ', '_')
            temp_filename = f"{safe_username}_{timestamp}.adi"
            temp_path = os.path.join(settings.MEDIA_ROOT, 'temp', temp_filename)

            # Создаем директорию если не существует
            os.makedirs(os.path.dirname(temp_path), exist_ok=True)

            with open(temp_path, 'w', encoding='utf-8') as f:
                f.write(converted_content)

            # URL для скачивания
            download_url = f'/media/temp/{temp_filename}'
            download_filename = temp_filename

            # Превью - первые 5 записей
            preview_lines = converted_content.split('\n')
            preview = '\n'.join(preview_lines[:20])  # Показываем первые 20 строк

            messages.success(request, f'Файл успешно конвертирован! Обработано записей: {len(converted_records)}')

        except Exception as e:
            messages.error(request, f'Ошибка при обработке файла: {str(e)}')
            return render(request, 'qo100/converter.html')

    return render(request, 'qo100/converter.html', {
        'preview': preview,
        'download_url': download_url,
        'download_filename': download_filename,
        'page_title': 'Конвертер лога QO-100',
        'page_subtitle': 'Конвертер ADIF файлов для спутника QO-100',
    })


def qo100_converter_download(request):
    """
    Скачивание конвертированного файла
    """
    import os
    from django.conf import settings
    from django.http import HttpResponse

    if not request.user.is_authenticated:
        return redirect('login_page')

    # Получаем параметр filename из URL
    filename = request.GET.get('filename', '')
    if not filename:
        messages.error(request, 'Файл не найден')
        return redirect('qo100_converter')

    # Проверяем безопасность пути
    safe_filename = os.path.basename(filename)
    file_path = os.path.join(settings.MEDIA_ROOT, 'temp', safe_filename)

    # Проверяем существование файла
    if not os.path.exists(file_path):
        messages.error(request, 'Файл не найден или срок хранения истёк')
        return redirect('qo100_converter')

    # Отдаем файл
    with open(file_path, 'rb') as f:
        content = f.read()

    response = HttpResponse(content, content_type='application/octet-stream')
    response['Content-Disposition'] = f'attachment; filename="{safe_filename}"'
    return response


def cosmos_diploma(request):
    """
    Страница заявки на диплом Космос
    """
    if not request.user.is_authenticated:
        return redirect('login_page')

    # Проверяем, не заблокирован ли пользователь
    is_blocked, reason = check_user_blocked(request.user)
    if is_blocked:
        return render(request, 'blocked.html', {'reason': reason})

    # Получаем данные пользователя из базы
    user = request.user
    profile = None
    try:
        profile = user.radio_profile
    except RadioProfile.DoesNotExist:
        pass

    # Формируем ФИО из базы
    full_name = ''
    if profile:
        if profile.first_name and profile.last_name:
            full_name = f"{profile.last_name} {profile.first_name}"
        elif profile.last_name:
            full_name = profile.last_name
        elif profile.first_name:
            full_name = profile.first_name
    if not full_name and (user.first_name or user.last_name):
        full_name = f"{user.last_name or ''} {user.first_name or ''}".strip()

    # Получаем email из базы
    email = user.email or ''

    # Получаем основной позывной (user.username)
    main_callsign = user.username.upper()

    # Получаем дополнительные позывные из профиля
    other_callsigns_list = []
    if profile and profile.my_callsigns:
        import json
        try:
            my_callsigns = profile.my_callsigns
            if isinstance(my_callsigns, str):
                my_callsigns = json.loads(my_callsigns)
            if isinstance(my_callsigns, list):
                for item in my_callsigns:
                    if isinstance(item, dict):
                        if item.get('name'):
                            other_callsigns_list.append(item['name'].upper())
                    elif isinstance(item, str):
                        other_callsigns_list.append(item.upper())
        except (json.JSONDecodeError, TypeError):
            pass

    # Удаляем основной позывной из списка дополнительных, если он там есть
    other_callsigns_list = [c for c in other_callsigns_list if c != main_callsign]

    # Загружаем данные из куки для предзаполнения формы
    def get_cookie(name, default=''):
        return request.COOKIES.get(name, default) or default

    # Инициализируем данные формы
    form_data = {
        'main_callsign': get_cookie('cosmos_main_callsign', main_callsign),
        'full_name': get_cookie('cosmos_full_name', full_name),
        'email': get_cookie('cosmos_email', email),
        'phone': get_cookie('cosmos_phone', ''),
        'info': get_cookie('cosmos_info', ''),
    }

    # Загружаем дополнительные позывные из куки
    other_callsigns_cookie = get_cookie('cosmos_other_callsigns', '')
    if other_callsigns_cookie:
        try:
            saved_callsigns = json.loads(other_callsigns_cookie)
            if isinstance(saved_callsigns, list):
                other_callsigns_list = saved_callsigns
        except json.JSONDecodeError:
            pass

    if request.method == 'POST':
        try:
            import json
            import os
            from datetime import datetime
            from django.conf import settings
            from docx import Document
            from docx.shared import Pt
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls

            # Получаем данные из формы
            main_callsign = request.POST.get('main_callsign', '').strip().upper()
            full_name = request.POST.get('full_name', '').strip()
            email = request.POST.get('email', '').strip()
            phone = request.POST.get('phone', '').strip()
            info = request.POST.get('info', '').strip()

            # Получаем дополнительные позывные (поддерживаем оба формата)
            # Новый формат: массив из input полей
            other_callsigns_from_inputs = request.POST.getlist('other_callsigns_names[]')
            # Старый формат: JSON из hidden field
            other_callsigns_raw = request.POST.get('other_callsigns_json', '[]')

            # Используем данные из input полей если есть, иначе из JSON
            if other_callsigns_from_inputs:
                other_callsigns = [c.strip().upper() for c in other_callsigns_from_inputs if c.strip()]
            else:
                try:
                    other_callsigns = json.loads(other_callsigns_raw)
                    if isinstance(other_callsigns, list):
                        other_callsigns = [c.upper() if isinstance(c, str) else c for c in other_callsigns]
                except (json.JSONDecodeError, TypeError):
                    other_callsigns = []

            # Валидация
            if not main_callsign:
                messages.error(request, 'Позывной обязателен')
                return render(request, 'cosmos_diploma.html', {
                    'form_data': form_data,
                    'other_callsigns': other_callsigns_list,
                })

            if not full_name:
                messages.error(request, 'ФИО обязательно')
                return render(request, 'cosmos_diploma.html', {
                    'form_data': form_data,
                    'other_callsigns': other_callsigns_list,
                })

            if not email:
                messages.error(request, 'Email обязателен')
                return render(request, 'cosmos_diploma.html', {
                    'form_data': form_data,
                    'other_callsigns': other_callsigns_list,
                })

            # Формируем список всех позывных для поиска QSO
            all_callsigns = [main_callsign] + [c for c in other_callsigns if c and c != main_callsign]

            # Запрос QSO из базы данных - ищем по всем позывным
            qso_data = []
            try:
                from django.db import connection

                # Создаем условие WHERE с учетом всех позывных и user_id
                if all_callsigns:
                    placeholders = ','.join(['%s'] * len(all_callsigns))
                    query = f"""
                        SELECT * FROM (
                            SELECT DISTINCT ON (callsign)
                                LEFT(COALESCE(gridsquare, ''), 4) as gridsquare,
                                date, band, time, mode, callsign,
                                rst_sent, rst_rcvd
                            FROM tlog_qso
                            WHERE user_id = %s AND my_callsign IN ({placeholders}) AND prop_mode = 'SAT'
                            ORDER BY callsign, date, time
                        ) AS distinct_qsos
                        ORDER BY date, time
                    """
                    with connection.cursor() as cursor:
                        # Параметры: user_id + все позывные
                        cursor.execute(query, [user.id] + all_callsigns)
                        qso_data = cursor.fetchall()
            except Exception as e:
                messages.warning(request, f'Не удалось получить данные QSO: {str(e)}')
                qso_data = []

            # Проверяем минимальное количество записей
            if len(qso_data) < 100:
                pass  # Предупреждение убрано - сообщение будет показано на странице загрузки файла

            # Путь к шаблону
            template_path = os.path.join(settings.BASE_DIR, 'tlog', 'cosmos', 'Cosmos.docx')

            if not os.path.exists(template_path):
                messages.error(request, 'Шаблон документа не найден')
                return render(request, 'cosmos_diploma.html', {
                    'form_data': form_data,
                    'other_callsigns': other_callsigns_list,
                })

            # Открываем документ
            doc = Document(template_path)

            # Форматируем данные
            def format_date(date_val):
                if date_val is None:
                    return ""
                if hasattr(date_val, 'strftime'):
                    return date_val.strftime("%d.%m.%Y")
                return str(date_val)

            def format_time(time_val):
                if time_val is None:
                    return ""
                if hasattr(time_val, 'strftime'):
                    return time_val.strftime("%H:%M")
                return str(time_val)

            def set_cell_border(cell):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                borders = parse_xml(
                    "<w:tcBorders {}><w:top w:val=\"single\" w:sz=\"2\"/><w:left w:val=\"single\" w:sz=\"2\"/><w:right w:val=\"single\" w:sz=\"2\"/><w:bottom w:val=\"single\" w:sz=\"2\"/><w:insideH w:val=\"single\" w:sz=\"2\"/><w:insideV w:val=\"single\" w:sz=\"2\"/></w:tcBorders>".format(
                        nsdecls("w")))
                tcPr.append(borders)

            def set_cell_shading(cell, color):
                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
                cell._tc.get_or_add_tcPr().append(shading_elm)

            def set_col_width(table, col_idx, width):
                for row in table.rows:
                    row.cells[col_idx].width = width

            # Формируем строку с другими позывными
            other_callsigns_str = ', '.join(other_callsigns) if other_callsigns else ''

            # Подсчеты
            count25 = (len(qso_data) // 25) * 25
            current_date = datetime.now().strftime("%d.%m.%Y")
            first_word_name = full_name.split()[0] if full_name else ""

            # Заполняем документ
            for paragraph in doc.paragraphs:
                if "{{YOUR_NAME}}" in paragraph.text:
                    parts = paragraph.text.split("{{YOUR_NAME}}")
                    paragraph.clear()
                    for part in parts[:-1]:
                        paragraph.add_run(part)
                        run = paragraph.add_run(full_name)
                        run.bold = True
                        run.font.size = Pt(14)
                    paragraph.add_run(parts[-1])
                elif "{{MY_CALLSIGN}}" in paragraph.text:
                    parts = paragraph.text.split("{{MY_CALLSIGN}}")
                    paragraph.clear()
                    for part in parts[:-1]:
                        paragraph.add_run(part)
                        run = paragraph.add_run(main_callsign)
                        run.bold = True
                        run.font.size = Pt(14)
                    paragraph.add_run(parts[-1])
                elif "{{OTHER_CALLSIGNS}}" in paragraph.text:
                    parts = paragraph.text.split("{{OTHER_CALLSIGNS}}")
                    paragraph.clear()
                    for part in parts[:-1]:
                        paragraph.add_run(part)
                        run = paragraph.add_run(other_callsigns_str)
                        run.bold = True
                        run.font.size = Pt(14)
                    paragraph.add_run(parts[-1])
                elif "{{YOUR_EMAIL}}" in paragraph.text:
                    parts = paragraph.text.split("{{YOUR_EMAIL}}")
                    paragraph.clear()
                    for part in parts[:-1]:
                        paragraph.add_run(part)
                        run = paragraph.add_run(email)
                        run.bold = True
                        run.font.size = Pt(14)
                    paragraph.add_run(parts[-1])
                elif "{{YOUR_PHONE}}" in paragraph.text:
                    parts = paragraph.text.split("{{YOUR_PHONE}}")
                    paragraph.clear()
                    for part in parts[:-1]:
                        paragraph.add_run(part)
                        run = paragraph.add_run(phone)
                        run.bold = True
                        run.font.size = Pt(14)
                    paragraph.add_run(parts[-1])
                elif "{{INITIAL}}" in paragraph.text:
                    parts = paragraph.text.split("{{INITIAL}}")
                    paragraph.clear()
                    for part in parts[:-1]:
                        paragraph.add_run(part)
                        run = paragraph.add_run(str(len(qso_data)))
                        run.bold = True
                        run.font.size = Pt(14)
                    paragraph.add_run(parts[-1])
                elif "{{INFO}}" in paragraph.text:
                    parts = paragraph.text.split("{{INFO}}")
                    paragraph.clear()
                    for part in parts[:-1]:
                        paragraph.add_run(part)
                        run = paragraph.add_run(info)
                        run.bold = True
                        run.font.size = Pt(14)
                    paragraph.add_run(parts[-1])
                elif "{{BOTTOM_LINE}}" in paragraph.text:
                    paragraph.clear()
                    run = paragraph.add_run("Подпись ")
                    run.font.size = Pt(14)
                    run = paragraph.add_run(first_word_name)
                    run.bold = True
                    run.font.size = Pt(14)
                    run = paragraph.add_run("       Позывной ")
                    run.font.size = Pt(14)
                    run = paragraph.add_run(main_callsign)
                    run.bold = True
                    run.font.size = Pt(14)
                    run = paragraph.add_run("          Дата ")
                    run.font.size = Pt(14)
                    run = paragraph.add_run(current_date)
                    run.bold = True
                    run.font.size = Pt(14)
                elif "{{count25}}" in paragraph.text:
                    original_run = paragraph.runs[0] if paragraph.runs else None
                    original_font_size = original_run.font.size if original_run else None
                    original_font_name = original_run.font.name if original_run else None
                    original_bold = original_run.bold if original_run else None

                    parts = paragraph.text.split("{{count25}}")
                    paragraph.clear()
                    for part in parts[:-1]:
                        paragraph.add_run(part)
                        if original_font_size:
                            paragraph.runs[-1].font.size = original_font_size
                        if original_font_name:
                            paragraph.runs[-1].font.name = original_font_name
                        if original_bold:
                            paragraph.runs[-1].bold = True
                    run = paragraph.add_run(str(count25))
                    run.bold = True
                    run.font.size = Pt(22)
                    paragraph.add_run(parts[-1])
                    if original_font_size:
                        paragraph.runs[-1].font.size = original_font_size
                    if original_font_name:
                        paragraph.runs[-1].font.name = original_font_name
                    if original_bold:
                        paragraph.runs[-1].bold = True
                elif "{{TABLE}}" in paragraph.text:
                    paragraph.clear()
                    if qso_data:
                        table = doc.add_table(rows=1, cols=9)
                        headers = ["№", "QTH-loc", "Дата", "Диапазон", "Время", "Модуляция", "Позывной", "Рапорт переданный", "Рапорт принятый"]
                        header_cells = table.rows[0].cells
                        for i, header in enumerate(headers):
                            header_cells[i].text = header
                            run = header_cells[i].paragraphs[0].runs[0]
                            run.bold = True
                            run.font.size = Pt(10)
                            set_cell_border(header_cells[i])
                            set_cell_shading(header_cells[i], "D9D9D9")
                        for row_num, row_data in enumerate(qso_data, 1):
                            row = table.add_row().cells
                            row[0].text = str(row_num)
                            row[1].text = row_data[0] or ""
                            row[2].text = format_date(row_data[1])
                            row[3].text = row_data[2] or ""
                            row[4].text = format_time(row_data[3])
                            row[5].text = row_data[4] or ""
                            row[6].text = row_data[5] or ""
                            row[7].text = row_data[6] or ""
                            row[8].text = row_data[7] or ""
                            for cell in row:
                                cell.paragraphs[0].runs[0].font.size = Pt(10)
                                set_cell_border(cell)
                        set_col_width(table, 0, Pt(10))

            # Формируем имя файла
            safe_username = user.username.replace(' ', '_')
            output_filename = f"Заявка_Cosmos_{safe_username}.docx"

            # Сохраняем во временный буфер
            from io import BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Сохраняем файл во временный файл для скачивания
            import tempfile
            import os
            from django.conf import settings

            temp_dir = os.path.join(settings.BASE_DIR, 'temp')
            os.makedirs(temp_dir, exist_ok=True)
            temp_file_path = os.path.join(temp_dir, f'cosmos_{user.id}_{int(time.time())}.docx')

            with open(temp_file_path, 'wb') as f:
                f.write(buffer.getvalue())

            # Сохраняем путь к файлу в сессии
            request.session['cosmos_download_path'] = temp_file_path
            request.session['cosmos_download_filename'] = output_filename

            # Сохраняем данные формы в сессии
            request.session['cosmos_form_data'] = {
                'main_callsign': main_callsign,
                'full_name': full_name,
                'email': email,
                'phone': phone,
                'info': info,
                'other_callsigns': other_callsigns,
            }

            # Формируем сообщение
            if len(qso_data) >= 100:
                messages.success(request, f'✅ Заявка сформирована! Найдено {len(qso_data)} уникальных QSO.')
            else:
                messages.warning(request, f'⚠️ Внимание! В заявке найдено только {len(qso_data)} уникальных QSO. Для диплома необходимо минимум 100.')

            # Редирект на страницу с сообщением
            return redirect('cosmos_diploma')

        except Exception as e:
            messages.error(request, f'Ошибка при формировании заявки: {str(e)}')
            return render(request, 'cosmos_diploma.html', {
                'form_data': form_data,
                'other_callsigns': other_callsigns_list,
            })

    # GET запрос - показываем форму
    # Проверяем, есть ли готовый файл для скачивания
    has_cosmos_file = 'cosmos_download_path' in request.session

    # Загружаем данные формы из сессии если есть
    form_data_session = request.session.get('cosmos_form_data')
    if form_data_session:
        form_data['main_callsign'] = form_data_session.get('main_callsign', main_callsign)
        form_data['full_name'] = form_data_session.get('full_name', full_name)
        form_data['email'] = form_data_session.get('email', email)
        form_data['phone'] = form_data_session.get('phone', '')
        form_data['info'] = form_data_session.get('info', '')
        # Загружаем дополнительные позывные из сессии
        other_callsigns_from_session = form_data_session.get('other_callsigns', [])
        if other_callsigns_from_session:
            other_callsigns_list = other_callsigns_from_session

    return render(request, 'cosmos_diploma.html', {
        'form_data': form_data,
        'other_callsigns': other_callsigns_list,
        'page_title': 'Заявка на диплом Cosmos',
        'page_subtitle': 'Формирование заявки на диплом Cosmos (СРР)',
        'has_cosmos_file': has_cosmos_file,
    })


def cosmos_download(request):
    """
    Скачивание сформированного файла заявки Cosmos
    """
    if not request.user.is_authenticated:
        return redirect('login_page')

    download_path = request.session.get('cosmos_download_path')
    download_filename = request.session.get('cosmos_download_filename')

    if not download_path or not os.path.exists(download_path):
        messages.error(request, 'Файл заявки не найден. Пожалуйста, сформируйте заявку заново.')
        return redirect('cosmos_diploma')

    try:
        with open(download_path, 'rb') as f:
            content = f.read()

        # Удаляем временный файл
        os.remove(download_path)

        # Очищаем сессию
        del request.session['cosmos_download_path']
        del request.session['cosmos_download_filename']

        response = HttpResponse(
            content,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        from django.utils.encoding import escape_uri_path
        response['Content-Disposition'] = f'attachment; filename="{escape_uri_path(download_filename)}"'
        return response

    except Exception as e:
        messages.error(request, f'Ошибка при скачивании файла: {str(e)}')
        return redirect('cosmos_diploma')